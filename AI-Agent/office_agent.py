import json
import os
import base64
import cv2
from io import BytesIO
from PIL import Image
from openai import OpenAI
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from faster_whisper import WhisperModel

# 1. 连接本地 Ollama (Gemma 4 Cloud)
client = OpenAI(
    base_url='http://localhost:11434/v1',
    api_key='ollama'
)

# ----------------- 本地工具实现 -----------------

def extract_slides_text(video_path, sample_interval_sec=10, diff_threshold=35.0, max_slides=6):
    """
    智能截取录屏中的 PPT 幻灯片，并通过 Gemma 4 原生多模态视觉能力读取画面重点
    """
    if not os.path.exists(video_path):
        return f"错误：找不到视频文件 {video_path}"
    
    print(f"\n[🖼️  正在分析录屏 PPT 画面...] 正在扫描视频切换帧: {video_path}")
    try:
        cap = cv2.VideoCapture(video_path)
        fps = cap.get(cv2.CAP_PROP_FPS)
        if fps == 0:
            return "无法读取视频帧率。"
        
        frame_interval = int(fps * sample_interval_sec)
        prev_gray = None
        slide_summaries = []
        slide_count = 0
        frame_idx = 0
        
        while cap.isOpened() and slide_count < max_slides:
            ret, frame = cap.read()
            if not ret:
                break
            
            if frame_idx % frame_interval == 0:
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                is_new_slide = False
                
                if prev_gray is None:
                    is_new_slide = True
                    prev_gray = gray
                else:
                    diff = cv2.absdiff(prev_gray, gray)
                    if diff.mean() > diff_threshold:
                        is_new_slide = True
                        prev_gray = gray
                
                if is_new_slide:
                    slide_count += 1
                    print(f"  └─ 捕捉到第 {slide_count} 页 PPT 幻灯片，正在使用 Gemma 4 视觉分析...")
                    
                    # 缩放图像以加快传输与推理
                    resized_frame = cv2.resize(frame, (1024, int(1024 * frame.shape[0] / frame.shape[1])))
                    _, buffer = cv2.imencode('.jpg', resized_frame)
                    base64_image = base64.b64encode(buffer).decode('utf-8')
                    
                    # 直接调用 Gemma 4 多模态视觉接口
                    try:
                        vision_resp = client.chat.completions.create(
                            model="gemma4:31b-cloud",
                            messages=[{
                                "role": "user",
                                "content": [
                                    {"type": "text", "text": "请简要提取这张会议 PPT 幻灯片中的核心文字标题、数据与关键要点："},
                                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                                ]
                            }]
                        )
                        slide_text = vision_resp.choices[0].message.content
                        slide_summaries.append(f"--- [幻灯片第 {slide_count} 页要点] ---\n{slide_text}")
                    except Exception as ve:
                        print(f"     视觉分析略过: {str(ve)}")
                        
            frame_idx += 1
        
        cap.release()
        total_extracted = "\n\n".join(slide_summaries)
        print(f"[✅ PPT 视觉提取完成] 共分析 {len(slide_summaries)} 个关键幻灯片画面。")
        return total_extracted if total_extracted else "录屏中未检测到明显的 PPT 幻灯片切换。"
    except Exception as e:
        print(f"[⚠️ PPT 画面提取异常] {str(e)}，将仅依靠语音逐字稿生成公文。")
        return "未能提取 PPT 画面，请完全依据语音逐字稿内容进行整理。"

def transcribe_audio(file_path):
    if not os.path.exists(file_path):
        return f"错误：找不到文件 {file_path}"
    
    print(f"\n[🎙️  正在转录音视频语音...] 请稍候: {file_path}")
    try:
        model = WhisperModel("base", device="cpu", compute_type="int8")
        segments, info = model.transcribe(file_path, beam_size=5)
        
        transcript = [segment.text for segment in segments]
        full_transcript = "".join(transcript).strip()
        print(f"[✅ 语音转录完成] 识别语言: {info.language}，总字数: {len(full_transcript)}")
        
        txt_path = os.path.splitext(file_path)[0] + "_逐字稿.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(full_transcript)
            
        return f"语音转录成功！以下为转录文本内容：\n{full_transcript}"
    except Exception as e:
        return f"语音转录失败，错误原因: {str(e)}"

def generate_peiching_minit_curai(
    file_path,
    tajuk_program,
    tarikh,
    masa,
    tempat,
    penganjur,
    penceramah,
    salinan_kepada="SEMUA GURU",
    pengisian_items=[],
    jawatan_guru="Guru Penolong",
    guru_besar_nama="CHEN LEE LEE"
):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    header_p = doc.add_paragraph()
    header_run = header_p.add_run("SJK(C) PEI CHING,\n32700 BERUAS,\nPERAK\n")
    header_run.bold = True
    header_run.font.size = Pt(11)
    
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("KERTAS MINIT CURAI")
    title_run.bold = True
    title_run.font.size = Pt(13)
    title_run.font.underline = True
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    meta_info = [
        ("Kepada", "GURU BESAR"),
        ("Daripada", "PANG SHENG HONG"),
        ("Salinan kepada", salinan_kepada),
        ("Tajuk Program", str(tajuk_program).upper()),
        ("Tarikh", str(tarikh).upper()),
        ("Masa", str(masa)),
        ("Tempat", str(tempat).upper()),
        ("Penganjur", str(penganjur).upper()),
        ("Penceramah", str(penceramah).upper()),
    ]
    
    for label, val in meta_info:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        run_label = p.add_run(f"{label:<16}: ")
        run_label.bold = True
        run_val = p.add_run(val)
        
    content_head = doc.add_paragraph()
    content_head.paragraph_format.space_before = Pt(8)
    head_run = content_head.add_run("Ringkasan Kandungan Pengisian:")
    head_run.bold = True
    head_run.font.underline = True
    
    p_main = doc.add_paragraph()
    p_main_run = p_main.add_run(f"1. Bengkel / Taklimat daripada {str(penceramah).upper()}:")
    p_main_run.bold = True
    
    for item in pengisian_items:
        sub_title = item.get("sub_title", "")
        details = item.get("details", [])
        
        if sub_title:
            sub_p = doc.add_paragraph()
            sub_p.paragraph_format.left_indent = Inches(0.3)
            sub_p.paragraph_format.space_after = Pt(2)
            sub_run = sub_p.add_run(sub_title)
            sub_run.bold = True
            
        for detail in details:
            d_p = doc.add_paragraph(style='List Bullet')
            d_p.paragraph_format.left_indent = Inches(0.5)
            d_p.paragraph_format.space_after = Pt(2)
            d_p.add_run(detail)
            
    doc.add_paragraph().paragraph_format.space_before = Pt(12)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell_left = table.cell(0, 0)
    cell_left.width = Inches(3.2)
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.line_spacing = 1.15
    p_left.add_run("Disediakan oleh,\n\n..................................................\n")
    p_left.add_run("(PANG SHENG HONG)\n").bold = True
    p_left.add_run(f"{jawatan_guru}\nSJKC Pei Ching\n")
    p_left.add_run(f"Tarikh : {tarikh}")
    
    cell_right = table.cell(0, 1)
    cell_right.width = Inches(3.2)
    p_right = cell_right.paragraphs[0]
    p_right.paragraph_format.line_spacing = 1.15
    p_right.add_run("Disahkan oleh,\n\n..................................................\n")
    if guru_besar_nama:
        p_right.add_run(f"({guru_besar_nama.upper()})\n").bold = True
    else:
        p_right.add_run("(                                             )\n").bold = True
    p_right.add_run("Guru Besar\nSJKC Pei Ching\n")
    p_right.add_run("Tarikh : ")
    
    doc.save(file_path)
    return f"成功按照培青华小标准格式生成公文：{file_path}"

# ----------------- 工具 Schema 定义 -----------------

tools = [
    {
        "type": "function",
        "function": {
            "name": "transcribe_audio",
            "description": "读取音频(.mp3, .m4a)或录屏(.mp4)的语音内容并转录为文本逐字稿",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "音视频文件路径"}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "extract_slides_text",
            "description": "自动截取录屏(.mp4)中播放的 PPT 幻灯片并使用视觉模型读取画面文字，用于补充关键数据",
            "parameters": {
                "type": "object",
                "properties": {
                    "video_path": {"type": "string", "description": "录屏文件路径"}
                },
                "required": ["video_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_peiching_minit_curai",
            "description": "严格按照 SJK(C) PEI CHING 培青华小公文标准生成规范的 Minit Curai Word 文档",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "输出的 docx 文件名，例如 'Minit_Curai.docx'"},
                    "tajuk_program": {"type": "string", "description": "会议/简报完整名称"},
                    "tarikh": {"type": "string", "description": "会议日期，例如 '7 JULAI 2026'"},
                    "masa": {"type": "string", "description": "会议时间，例如 '4:00 p.m.'"},
                    "tempat": {"type": "string", "description": "地点或平台，例如 'SECARA ATAS TALIAN (WEBEX)'"},
                    "penganjur": {"type": "string", "description": "主办单位"},
                    "penceramah": {"type": "string", "description": "主讲人姓名"},
                    "salinan_kepada": {"type": "string", "description": "抄送对象，默认为 'SEMUA GURU'"},
                    "pengisian_items": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "sub_title": {"type": "string", "description": "子项目标题，如 'i. Pengenalan Program'"},
                                "details": {"type": "array", "items": {"type": "string"}, "description": "具体要点与决策列表"}
                            },
                            "required": ["sub_title", "details"]
                        },
                        "description": "按 i, ii, iii 组织的会议内容与决议结构"
                    },
                    "jawatan_guru": {"type": "string", "description": "职衔，默认为 'Guru Penolong'"}
                },
                "required": ["file_path", "tajuk_program", "tarikh", "masa", "tempat", "penganjur", "penceramah", "pengisian_items"]
            }
        }
    }
]

# ----------------- 交互循环 -----------------

def interactive_agent():
    print("="*60)
    print("🤖 培青华小专属 Office & 多模态 AI Agent (无依赖升级版) 已就绪！")
    print("="*60)
    
    messages = [
        {"role": "system", "content": """
        你是由潘盛鸿 (Pang Sheng Hong) 老师打造的专属教学与公文助理。
        工作单位：SJK(C) PEI CHING, 32700 BERUAS, PERAK (霹雳木威培青华小)。
        
        公文处理铁律：
        1. 当用户提供录屏（.mp4）时，你可以调用 transcribe_audio 获取语音，并可调用 extract_slides_text 捕获 PPT 画面文字，结合两者生成最详实的会议内容。
        2. 生成 Minit Curai 时，必须调用 generate_peiching_minit_curai 工具，确保输出带有标准校头、9项元数据、i/ii 层级正文以及双栏签署表。
        3. 准备人固定为 Pang Sheng Hong，审核人默认为 CHEN LEE LEE 校长。
        """}
    ]
    
    while True:
        user_input = input("\n你: ")
        if user_input.lower() in ['exit', 'quit', '退出']:
            print("Agent 已安全退出。")
            break
        if not user_input.strip():
            continue
            
        messages.append({"role": "user", "content": user_input})
        
        while True:
            response = client.chat.completions.create(
                model="gemma4:31b-cloud",
                messages=messages,
                tools=tools
            )
            
            message = response.choices[0].message
            messages.append(message)
            
            if message.tool_calls:
                for tool_call in message.tool_calls:
                    func_name = tool_call.function.name
                    args = json.loads(tool_call.function.arguments)
                    print(f"\n[⚙️ 触发工具] {func_name} | 参数: {args}")
                    
                    if func_name == "transcribe_audio":
                        result = transcribe_audio(args.get("file_path"))
                    elif func_name == "extract_slides_text":
                        result = extract_slides_text(args.get("video_path"))
                    elif func_name == "generate_peiching_minit_curai":
                        result = generate_peiching_minit_curai(
                            file_path=args.get("file_path"),
                            tajuk_program=args.get("tajuk_program"),
                            tarikh=args.get("tarikh"),
                            masa=args.get("masa"),
                            tempat=args.get("tempat"),
                            penganjur=args.get("penganjur", args.get("penceramah")),
                            penceramah=args.get("penceramah"),
                            salinan_kepada=args.get("salinan_kepada", "SEMUA GURU"),
                            pengisian_items=args.get("pengisian_items", []),
                            jawatan_guru=args.get("jawatan_guru", "Guru Penolong")
                        )
                    else:
                        result = "未知工具"
                    
                    messages.append({
                        "role": "tool",
                        "tool_call_id": tool_call.id,
                        "content": str(result)
                    })
            else:
                print(f"\nAgent: {message.content}")
                break

if __name__ == "__main__":
    interactive_agent()