import json
import os
import base64
import cv2
import tempfile
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from faster_whisper import WhisperModel
import streamlit as st
from openai import OpenAI

# ================= 设置页面信息 =================
st.set_page_config(page_title="马来西亚教师专属 AI 助理", page_icon="🤖", layout="wide")
st.title("🤖 教师专属 Office & 多模态 AI Agent")
st.markdown("通过上传会议录屏或音频，AI 能够自动提取语音并分析画面内容，最终生成标准格式的 Minit Curai。")

# ================= 侧边栏配置 =================
with st.sidebar:
    st.header("⚙️ 配置参数")
    # 尝试从 Streamlit Secrets 中读取 API Key
    default_api_key = st.secrets.get("GOOGLE_API_KEY", "") if hasattr(st, "secrets") and "GOOGLE_API_KEY" in st.secrets else ""
    
    if not default_api_key:
        api_key = st.text_input("API Key (必填)", value="", type="password", placeholder="请填入你的 Google AI Studio API Key")
    else:
        api_key = default_api_key
        st.success("✅ 已自动加载系统内置的 API Key")
        
    # 2026年最新 Google API 兼容 OpenAI 格式的地址
    base_url = st.text_input("Base URL", value="https://generativelanguage.googleapis.com/v1beta/openai/")
    model_name = st.text_input("模型名称", value="gemma-4-31b-it")
    
    st.divider()
    st.header("📂 上传文件")
    st.warning("⚠️ **云端系统限制**\n\n为确保全国教师的使用体验，系统已限制最大上传文件为 **500MB**。")
    
    with st.expander("🎥 视频文件超过 500MB 怎么办？(附教程)", expanded=False):
        st.markdown("""
        **强烈建议：使用免费且无需安装的浏览器压缩工具！**
        
        👉 **点击进入：[Compress.lol 极速视频压缩](https://compress.lol/)**
        
        **操作步骤：**
        1. 点击上方链接打开压缩网站。
        2. 把你录制的大视频拖进去。
        3. 等待浏览器在本地将其压缩到 500MB 以内。
        4. 下载压缩后的视频，重新上传到本系统即可！
        """)
        # 自动播放的动图演示（使用 WebP 格式体积更小）
        st.image(os.path.join(os.path.dirname(__file__), "tutorial.webp"), caption="操作演示")

    uploaded_file = st.file_uploader("上传录音/录屏 (mp4, mp3, m4a)", type=["mp4", "mp3", "m4a"])
    
    # 将上传的文件保存到本地临时路径供 cv2 和 whisper 读取
    current_file_path = None
    if uploaded_file is not None:
        # 使用 tempfile 保存
        temp_dir = tempfile.gettempdir()
        current_file_path = os.path.join(temp_dir, uploaded_file.name)
        with open(current_file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        st.success(f"文件已就绪: {uploaded_file.name}")
        # 提供给 Agent 一个系统提示，告诉它当前文件的路径
        st.session_state["current_file_path"] = current_file_path
        
    # 显示生成的供下载的文件
    if "generated_files" in st.session_state and st.session_state.generated_files:
        st.divider()
        st.header("📥 下载生成的公文")
        for fname, fbytes in st.session_state.generated_files.items():
            st.download_button(
                label=f"下载 {fname}", 
                data=fbytes, 
                file_name=fname, 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ================= 初始化 OpenAI 客户端 =================
client = OpenAI(
    base_url=base_url,
    api_key=api_key
)

# ================= 本地工具实现 (复用原代码) =================
def extract_slides_text(video_path, sample_interval_sec=10, diff_threshold=35.0, max_slides=6):
    if not os.path.exists(video_path):
        return f"错误：找不到视频文件 {video_path}"
    
    st.toast("🖼️ 正在分析录屏 PPT 画面...", icon="🔍")
    try:
        cap = cv2.VideoCapture(video_path)
        fps = cap.get(cv2.CAP_PROP_FPS)
        if fps == 0:
            return "无法读取视频帧率。"
        
        frame_interval = int(fps * sample_interval_sec)
        prev_gray = None
        slide_summaries = []
        slide_count = 0
        frame_idx = 0
        
        while cap.isOpened() and slide_count < max_slides:
            ret, frame = cap.read()
            if not ret:
                break
            
            if frame_idx % frame_interval == 0:
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                is_new_slide = False
                
                if prev_gray is None:
                    is_new_slide = True
                    prev_gray = gray
                else:
                    diff = cv2.absdiff(prev_gray, gray)
                    if diff.mean() > diff_threshold:
                        is_new_slide = True
                        prev_gray = gray
                
                if is_new_slide:
                    slide_count += 1
                    
                    resized_frame = cv2.resize(frame, (1024, int(1024 * frame.shape[0] / frame.shape[1])))
                    _, buffer = cv2.imencode('.jpg', resized_frame)
                    base64_image = base64.b64encode(buffer).decode('utf-8')
                    
                    try:
                        vision_resp = client.chat.completions.create(
                            model=model_name,
                            messages=[{
                                "role": "user",
                                "content": [
                                    {"type": "text", "text": "请简要提取这张会议 PPT 幻灯片中的核心文字标题、数据与关键要点："},
                                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                                ]
                            }]
                        )
                        slide_text = vision_resp.choices[0].message.content
                        slide_summaries.append(f"--- [幻灯片第 {slide_count} 页要点] ---\n{slide_text}")
                    except Exception as ve:
                        st.error(f"视觉分析略过: {str(ve)}")
                        
            frame_idx += 1
        
        cap.release()
        total_extracted = "\n\n".join(slide_summaries)
        return total_extracted if total_extracted else "录屏中未检测到明显的 PPT 幻灯片切换。"
    except Exception as e:
        return "未能提取 PPT 画面，请完全依据语音逐字稿内容进行整理。"

def transcribe_audio(file_path):
    if not os.path.exists(file_path):
        return f"错误：找不到文件 {file_path}"
    
    st.toast("🎙️ 正在转录音视频语音...", icon="⏳")
    try:
        model = WhisperModel("base", device="cpu", compute_type="int8")
        segments, info = model.transcribe(file_path, beam_size=5)
        
        transcript = [segment.text for segment in segments]
        full_transcript = "".join(transcript).strip()
        
        txt_path = os.path.splitext(file_path)[0] + "_逐字稿.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(full_transcript)
            
        return f"语音转录成功！以下为转录文本内容：\n{full_transcript}"
    except Exception as e:
        return f"语音转录失败，错误原因: {str(e)}"

def generate_minit_curai(
    file_path, tajuk_program, tarikh, masa, tempat, penganjur, penceramah,
    nama_sekolah="NAMA SEKOLAH", alamat_sekolah="ALAMAT SEKOLAH",
    nama_penyedia="NAMA GURU", jawatan_penyedia="Guru Penolong",
    nama_pengesah="NAMA GURU BESAR", jawatan_pengesah="Guru Besar",
    salinan_kepada="SEMUA GURU", pengisian_items=[]
):
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    header_p = doc.add_paragraph()
    header_run = header_p.add_run(f"{str(nama_sekolah).upper()},\n{str(alamat_sekolah).upper()}\n")
    header_run.bold = True
    header_run.font.size = Pt(11)
    
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("KERTAS MINIT CURAI")
    title_run.bold = True
    title_run.font.size = Pt(13)
    title_run.font.underline = True
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    meta_info = [
        ("Kepada", "GURU BESAR"),
        ("Daripada", str(nama_penyedia).upper()),
        ("Salinan kepada", salinan_kepada),
        ("Tajuk Program", str(tajuk_program).upper()),
        ("Tarikh", str(tarikh).upper()),
        ("Masa", str(masa)),
        ("Tempat", str(tempat).upper()),
        ("Penganjur", str(penganjur).upper()),
        ("Penceramah", str(penceramah).upper()),
    ]
    
    for label, val in meta_info:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        run_label = p.add_run(f"{label:<16}: ")
        run_label.bold = True
        run_val = p.add_run(val)
        
    content_head = doc.add_paragraph()
    content_head.paragraph_format.space_before = Pt(8)
    head_run = content_head.add_run("Ringkasan Kandungan Pengisian:")
    head_run.bold = True
    head_run.font.underline = True
    
    p_main = doc.add_paragraph()
    p_main_run = p_main.add_run(f"1. Bengkel / Taklimat daripada {str(penceramah).upper()}:")
    p_main_run.bold = True
    
    for item in pengisian_items:
        sub_title = item.get("sub_title", "")
        details = item.get("details", [])
        
        if sub_title:
            sub_p = doc.add_paragraph()
            sub_p.paragraph_format.left_indent = Inches(0.3)
            sub_p.paragraph_format.space_after = Pt(2)
            sub_run = sub_p.add_run(sub_title)
            sub_run.bold = True
            
        for detail in details:
            d_p = doc.add_paragraph(style='List Bullet')
            d_p.paragraph_format.left_indent = Inches(0.5)
            d_p.paragraph_format.space_after = Pt(2)
            d_p.add_run(detail)
            
    doc.add_paragraph().paragraph_format.space_before = Pt(12)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell_left = table.cell(0, 0)
    cell_left.width = Inches(3.2)
    p_left = cell_left.paragraphs[0]
    p_left.paragraph_format.line_spacing = 1.15
    p_left.add_run("Disediakan oleh,\n\n..................................................\n")
    p_left.add_run(f"({str(nama_penyedia).upper()})\n").bold = True
    p_left.add_run(f"{jawatan_penyedia}\n{nama_sekolah}\n")
    p_left.add_run(f"Tarikh : {tarikh}")
    
    cell_right = table.cell(0, 1)
    cell_right.width = Inches(3.2)
    p_right = cell_right.paragraphs[0]
    p_right.paragraph_format.line_spacing = 1.15
    p_right.add_run("Disahkan oleh,\n\n..................................................\n")
    if nama_pengesah:
        p_right.add_run(f"({str(nama_pengesah).upper()})\n").bold = True
    else:
        p_right.add_run("(                                             )\n").bold = True
    p_right.add_run(f"{jawatan_pengesah}\n{nama_sekolah}\n")
    p_right.add_run("Tarikh : ")
    
    doc.save(file_path)
    return "成功生成公文！【重要指示】请在回复中明确告诉用户：公文已成功生成，请点击网页左侧边栏底部的『📥 下载生成的公文』按钮进行下载。绝对不要在回复中提供虚假的文件下载链接！"

# ================= 工具 Schema 定义 =================
tools = [
    {
        "type": "function",
        "function": {
            "name": "transcribe_audio",
            "description": "读取音频(.mp3, .m4a)或录屏(.mp4)的语音内容并转录为文本逐字稿",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "音视频文件路径"}
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "extract_slides_text",
            "description": "自动截取录屏(.mp4)中播放的 PPT 幻灯片并使用视觉模型读取画面文字，用于补充关键数据",
            "parameters": {
                "type": "object",
                "properties": {
                    "video_path": {"type": "string", "description": "录屏文件路径"}
                },
                "required": ["video_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_minit_curai",
            "description": "严格按照标准格式生成规范的 Minit Curai Word 文档",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "输出的 docx 文件名，例如 'Minit_Curai.docx'"},
                    "tajuk_program": {"type": "string", "description": "会议/简报完整名称"},
                    "tarikh": {"type": "string", "description": "会议日期，例如 '7 JULAI 2026'"},
                    "masa": {"type": "string", "description": "会议时间，例如 '4:00 p.m.'"},
                    "tempat": {"type": "string", "description": "地点或平台，例如 'SECARA ATAS TALIAN (WEBEX)'"},
                    "penganjur": {"type": "string", "description": "主办单位"},
                    "penceramah": {"type": "string", "description": "主讲人姓名"},
                    "nama_sekolah": {"type": "string", "description": "学校名称"},
                    "alamat_sekolah": {"type": "string", "description": "学校地址"},
                    "nama_penyedia": {"type": "string", "description": "准备人姓名 (Disediakan oleh)"},
                    "jawatan_penyedia": {"type": "string", "description": "准备人职衔，默认 'Guru Penolong'"},
                    "nama_pengesah": {"type": "string", "description": "审核人姓名 (Disahkan oleh)"},
                    "jawatan_pengesah": {"type": "string", "description": "审核人职衔，默认 'Guru Besar'"},
                    "salinan_kepada": {"type": "string", "description": "抄送对象，默认为 'SEMUA GURU'"},
                    "pengisian_items": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "sub_title": {"type": "string", "description": "子项目标题，如 'i. Pengenalan Program'"},
                                "details": {"type": "array", "items": {"type": "string"}, "description": "具体要点与决策列表"}
                            },
                            "required": ["sub_title", "details"]
                        },
                        "description": "按 i, ii, iii 组织的会议内容与决议结构"
                    }
                },
                "required": ["file_path", "tajuk_program", "tarikh", "masa", "tempat", "penganjur", "penceramah", "pengisian_items", "nama_sekolah", "nama_penyedia", "nama_pengesah"]
            }
        }
    }
]

# ================= 交互核心 =================
system_prompt = """
你是马来西亚全国教师专属的教学与公文助理。

公文处理铁律：
1. 当用户提供录屏（.mp4）时，你可以调用 transcribe_audio 获取语音，并可调用 extract_slides_text 捕获 PPT 画面文字，结合两者生成最详实的会议内容。
2. 生成 Minit Curai 时，必须调用 generate_minit_curai 工具，确保输出带有标准校头、9项元数据、i/ii 层级正文以及双栏签署表。
3. 请主动向用户询问他们的学校名称、学校地址、姓名（作为准备人）、职位，以及校长的姓名（作为审核人），以便在生成的公文中正确填写这些信息。如果用户没有提供，你可以先用占位符（例如：NAMA SEKOLAH, NAMA GURU）生成。
"""

if "messages" not in st.session_state:
    st.session_state.messages = [{"role": "system", "content": system_prompt}]

# 显示历史消息 (过滤掉 system 提示和 tool 的原始调用，保持界面整洁)
for msg in st.session_state.messages:
    if msg["role"] == "user":
        with st.chat_message("user"):
            st.markdown(msg["content"])
    elif msg["role"] == "assistant" and msg.get("content"):
        with st.chat_message("assistant"):
            st.markdown(msg["content"])
    elif msg["role"] == "tool":
        # 可以选择显示工具执行结果
        with st.chat_message("assistant", avatar="⚙️"):
            st.info(f"✅ 工具执行完毕: {msg.get('name', 'Tool')}")

# 获取用户输入
if prompt := st.chat_input("输入你的指令，例如：'帮我整理刚才上传的会议视频并生成公文'"):
    
    # 如果用户上传了文件，我们悄悄在用户的提示词里附上文件路径
    if st.session_state.get("current_file_path"):
        context_prompt = f"{prompt}\n[系统提示：用户已上传文件，路径为: {st.session_state['current_file_path']}]"
    else:
        context_prompt = prompt

    st.session_state.messages.append({"role": "user", "content": context_prompt})
    
    with st.chat_message("user"):
        st.markdown(prompt)

    with st.chat_message("assistant"):
        message_placeholder = st.empty()
        
        with st.spinner("思考中..."):
            try:
                # 第一轮对话请求
                response = client.chat.completions.create(
                    model=model_name,
                    messages=st.session_state.messages,
                    tools=tools
                )
                
                response_message = response.choices[0].message
                # 需要把对象转成字典存储在 session_state 以免出现序列化问题，或者直接存对象
                st.session_state.messages.append(response_message.model_dump(exclude_none=True))

                # 如果模型决定调用工具
                if response_message.tool_calls:
                    for tool_call in response_message.tool_calls:
                        func_name = tool_call.function.name
                        args = json.loads(tool_call.function.arguments)
                        
                        with st.status(f"执行工具: {func_name}...", expanded=True):
                            st.write(f"参数: {args}")
                            
                            if func_name == "transcribe_audio":
                                result = transcribe_audio(args.get("file_path"))
                            elif func_name == "extract_slides_text":
                                result = extract_slides_text(args.get("video_path"))
                            elif func_name == "generate_minit_curai":
                                # 强制将文件保存在临时文件夹中，提取文件名
                                filename = os.path.basename(args.get("file_path", "Minit_Curai.docx"))
                                if not filename.endswith(".docx"):
                                    filename += ".docx"
                                save_path = os.path.join(tempfile.gettempdir(), filename)
                                
                                result = generate_minit_curai(
                                    file_path=save_path,
                                    tajuk_program=args.get("tajuk_program"),
                                    tarikh=args.get("tarikh"),
                                    masa=args.get("masa"),
                                    tempat=args.get("tempat"),
                                    penganjur=args.get("penganjur", args.get("penceramah")),
                                    penceramah=args.get("penceramah"),
                                    nama_sekolah=args.get("nama_sekolah", "NAMA SEKOLAH"),
                                    alamat_sekolah=args.get("alamat_sekolah", "ALAMAT SEKOLAH"),
                                    nama_penyedia=args.get("nama_penyedia", "NAMA GURU"),
                                    jawatan_penyedia=args.get("jawatan_penyedia", "Guru Penolong"),
                                    nama_pengesah=args.get("nama_pengesah", "NAMA GURU BESAR"),
                                    jawatan_pengesah=args.get("jawatan_pengesah", "Guru Besar"),
                                    salinan_kepada=args.get("salinan_kepada", "SEMUA GURU"),
                                    pengisian_items=args.get("pengisian_items", [])
                                )
                                
                                # 将生成的文件读取到 session_state 中，实现持久化下载
                                if os.path.exists(save_path):
                                    with open(save_path, "rb") as f:
                                        file_bytes = f.read()
                                    if "generated_files" not in st.session_state:
                                        st.session_state.generated_files = {}
                                    st.session_state.generated_files[filename] = file_bytes
                                    st.success(f"🎉 文件已生成并保存在临时目录，请在左侧侧边栏点击下载！")
                            else:
                                result = "未知工具"
                                
                            st.write("执行结果:", result)
                        
                        st.session_state.messages.append({
                            "role": "tool",
                            "tool_call_id": tool_call.id,
                            "name": func_name,
                            "content": str(result)
                        })
                        
                    # 把工具结果再发回给大模型，让它总结
                    with st.spinner("根据工具结果生成总结..."):
                        second_response = client.chat.completions.create(
                            model=model_name,
                            messages=st.session_state.messages
                        )
                        final_msg = second_response.choices[0].message
                        st.session_state.messages.append(final_msg.model_dump(exclude_none=True))
                        message_placeholder.markdown(final_msg.content)
                
                else:
                    # 如果没有调用工具，直接显示模型的回答
                    message_placeholder.markdown(response_message.content)
                    
            except Exception as e:
                st.error(f"请求出错: {str(e)}")
